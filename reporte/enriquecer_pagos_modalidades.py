"""Enriquece los DOCX formales con pagos/modalidades y arquitectura operacional vigente.

El bloque de pagos se deriva de ``outputs/analisis_pagos_modalidades.json`` y
mantiene explícitamente su naturaleza sintética. El mismo anexo documenta el
contrato Spark-native vigente para evitar que los Productos formales queden
anclados únicamente a la evidencia histórica ``local[*]``.

Los DOCX base son generados por la librería JS ``docx``. Para evitar depender de
nombres de estilos de python-docx, los headings y bordes se aplican en OOXML.
El enriquecimiento se aplica al Informe Técnico y a los siete Productos para que
ningún entregable formal quede rezagado respecto de la arquitectura operacional.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
EVIDENCIA = ROOT / "outputs/analisis_pagos_modalidades.json"
CHART_RATIO = ROOT / "outputs/charts/11_ratio_pago_contrato.png"
CHART_MODALIDADES = ROOT / "outputs/charts/12_modalidades_regimen.png"

DESTINOS = [
    ROOT / "reporte/Reporte_Tecnico_Prototipo_CGR_1.8.2.docx",
    *sorted((ROOT / "reporte/productos_formales").glob("Producto_*.docx")),
]

MARCADOR = "ANEXO TÉCNICO — PAGOS, MODALIDADES Y ARQUITECTURA OPERACIONAL"
REFERENCIAS = "Referencias Bibliográficas"


def _formatear_parrafo(parrafo, negrita=False):
    parrafo.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    parrafo.paragraph_format.space_after = Pt(3)
    for run in parrafo.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = negrita or run.bold


def _asignar_estilo_heading_ooxml(parrafo, nivel: int) -> None:
    p_pr = parrafo._p.get_or_add_pPr()
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        p_style = OxmlElement("w:pStyle")
        p_pr.insert(0, p_style)
    p_style.set(qn("w:val"), f"Heading{nivel}")


def _agregar_heading(doc, texto: str, nivel: int = 1):
    p = doc.add_paragraph()
    _asignar_estilo_heading_ooxml(p, nivel)
    run = p.add_run(texto)
    run.bold = True
    _formatear_parrafo(p, negrita=True)
    return p


def _agregar_texto(doc, texto: str):
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _formatear_parrafo(p)
    return p


def _aplicar_bordes_tabla(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "B7B7B7")


def _sombrear_celda(cell, fill="E7E6E6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _agregar_tabla(doc, pagos: dict, modalidades: dict):
    estados = pagos.get("estado_pago_contratos", {})
    clasif = modalidades.get("clasificacion_modalidad", {})
    filas = [
        ("Pagos sintéticos", pagos.get("payments_rows")),
        ("Contratos analizados", pagos.get("contracts_rows")),
        ("Pagos huérfanos", pagos.get("orphan_payments")),
        ("Contratos con pago parcial", estados.get("pago_parcial", 0)),
        ("Contratos pendientes sin pago", estados.get("pendiente_sin_pago", 0)),
        ("Señales de sobrepago a revisar", estados.get("sobrepago_senal_revisar", 0)),
        ("Demora P90 devengado→pagado (días)", pagos.get("dias_devengado_a_pagado_p90")),
        ("Modalidades especiales no inferibles solo por cuantía", clasif.get("especial_no_inferible_por_cuantia", 0)),
        ("Modalidades que requieren revisión de contexto", clasif.get("requiere_revision_contexto", 0)),
    ]
    table = doc.add_table(rows=1, cols=2)
    _aplicar_bordes_tabla(table)
    table.rows[0].cells[0].text = "Indicador"
    table.rows[0].cells[1].text = "Resultado"
    for cell in table.rows[0].cells:
        _sombrear_celda(cell)
        for p in cell.paragraphs:
            _formatear_parrafo(p, negrita=True)

    for etiqueta, valor in filas:
        cells = table.add_row().cells
        cells[0].text = str(etiqueta)
        cells[1].text = "N/D" if valor is None else str(valor)
        for cell in cells:
            for p in cell.paragraphs:
                _formatear_parrafo(p)
    return table


def _mover_bloque_antes_de_referencias(doc, marcador_parrafo) -> None:
    referencia = next((p for p in doc.paragraphs if p.text.strip() == REFERENCIAS), None)
    if referencia is None:
        raise ValueError(f"No se encontró el encabezado final {REFERENCIAS!r} en el DOCX.")

    nuevos = []
    elem = marcador_parrafo._p
    while elem is not None and not elem.tag.endswith("}sectPr"):
        nuevos.append(elem)
        elem = elem.getnext()
    if not nuevos:
        raise ValueError("No se detectó el bloque técnico recién creado.")
    for nuevo in nuevos:
        referencia._p.addprevious(nuevo)


def enriquecer(ruta: Path, evidencia: dict):
    if not ruta.exists():
        raise FileNotFoundError(ruta)
    doc = Document(ruta)
    texto_existente = "\n".join(p.text for p in doc.paragraphs)
    if MARCADOR in texto_existente:
        return False

    pagos = evidencia["payments"]
    modalidades = evidencia["modalidades"]

    marcador = _agregar_heading(doc, MARCADOR, 1)
    _agregar_texto(
        doc,
        "Este anexo cubre el análisis estadístico/exploratorio de pagos, montos contractuales y modalidades y actualiza la descripción de la arquitectura operacional. La evidencia de pagos usa datos sintéticos vinculados a contratos sintéticos; no representa información SIAF real ni determina la legalidad de una modalidad.",
    )
    _agregar_heading(doc, "Resultados reproducibles de pagos y modalidades", 2)
    _agregar_tabla(doc, pagos, modalidades)
    _agregar_texto(
        doc,
        "La comparación de modalidades frente a cuantía es únicamente referencial. Contratación Directa, Comparación de Precios, Subasta Inversa, acuerdos marco y otros supuestos pueden depender de condiciones que no se deducen del monto por sí solo; 'requiere revisión de contexto' no equivale a irregularidad.",
    )

    if CHART_RATIO.exists():
        _agregar_heading(doc, "Distribución de pagos", 2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(CHART_RATIO), width=Inches(6.1))
        _formatear_parrafo(p)
    if CHART_MODALIDADES.exists():
        _agregar_heading(doc, "Modalidades por régimen", 2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(CHART_MODALIDADES), width=Inches(6.1))
        _formatear_parrafo(p)

    _agregar_heading(doc, "Arquitectura operacional Spark-native", 2)
    _agregar_texto(
        doc,
        "La evidencia histórica de modelado Spark se conserva en modo local[*] para reproducibilidad. En paralelo, la ruta operacional TRAIN/INFERENCE admite un master Spark configurable. Cuando source.type=spark_sql, el conector devuelve un Spark DataFrame y el mapping, validación, preprocesamiento, feature engineering y MLlib permanecen en Spark sin toPandas(). Las medianas por objeto aprendidas en TRAIN se persisten como Parquet distribuido y los rankings de INFERENCE se escriben también como Parquet distribuido.",
    )
    _agregar_texto(
        doc,
        "Esta capacidad elimina el cuello de botella pandas del adaptador Spark anterior, pero no sustituye las pruebas de carga, particionamiento, seguridad, rendimiento y robustez que deben ejecutarse con el clúster y volumen institucionales reales.",
    )

    _agregar_heading(doc, "Limitación institucional", 2)
    _agregar_texto(
        doc,
        "El cierre literal con pagos SIAF/SEACE, tablas/vistas reales, ground truth, permisos, clúster, DEV/QA/PROD, SQL Server/SSRS, certificación y conformidad depende de la CGR y permanece registrado en el catálogo de dependencias institucionales.",
    )

    _mover_bloque_antes_de_referencias(doc, marcador)
    doc.save(ruta)
    return True


def main():
    evidencia = json.loads(EVIDENCIA.read_text(encoding="utf-8"))
    modificados = 0
    for ruta in DESTINOS:
        if enriquecer(ruta, evidencia):
            modificados += 1
            print(f"Enriquecido: {ruta.relative_to(ROOT)}")
        else:
            print(f"Ya contenía anexo técnico: {ruta.relative_to(ROOT)}")
    print(f"DOCX enriquecidos: {modificados}/{len(DESTINOS)}")


if __name__ == "__main__":
    main()
