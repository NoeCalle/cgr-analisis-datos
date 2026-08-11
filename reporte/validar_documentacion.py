#!/usr/bin/env python3
"""Auditoría estructural y editorial de los DOCX formales generados por el PoC."""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent
REPORT = ROOT / "Reporte_Tecnico_Prototipo_CGR_1.8.2.docx"
PRODUCTS = sorted((ROOT / "productos_formales").glob("Producto_*.docx"))
FILES = [REPORT] + PRODUCTS

# Estos tokens representan contratos obsoletos o narrativa interna de desarrollo
# que no debe volver a aparecer en los entregables vigentes.
BANNED = [
    "pct_no_competitiva",
    "S/. 400,000",
    "reentrenado automáticamente",
    "cumple_regla_fraccionamiento",
    "Sprint 2",
    "Sprint 3",
    "Sprint 4",
    "Etapa 2B",
    "Etapa 3B",
    "Etapa 4B",
    "Etapa 5B",
    "Correcciones P1",
    "conteo antiguo",
    "conteo anterior",
    "adaptador Spark anterior",
]

COMMON = [
    "Lista de Abreviaturas y Acrónimos",
    "Glosario",
    "Referencias Bibliográficas",
    "prototipo independiente",
]

PRODUCT_REPORT_HEADINGS = [
    "Resumen Ejecutivo",
    "1. Introducción",
    "2. Objetivo de Consultoría",
    "3. Productos Alcanzados",
    "4. Actividades Realizadas",
    "5. Grado de Cumplimiento del Producto",
    "6. Dificultades y Limitaciones Encontradas",
    "7. Conclusiones y Recomendaciones",
    "8. Anexos",
]

PLAN_HEADINGS = [
    "1. Introducción",
    "2. Objetivo de Consultoría",
    "3. Productos a Alcanzar",
    "4. Actividades a Cumplir por Cada Producto",
    "5. Cronograma",
    "6. Anexos",
]

FINAL_HEADINGS = [
    "Resumen Ejecutivo",
    "1. Introducción",
    "2. Objetivo de Consultoría",
    "3. Productos Alcanzados",
    "4. Conclusiones y Recomendaciones",
    "5. Anexos",
]


def all_text(doc: Document) -> str:
    partes = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            partes.append(" | ".join(cell.text for cell in row.cells))
    for section in doc.sections:
        for p in section.header.paragraphs:
            partes.append(p.text)
        for p in section.footer.paragraphs:
            partes.append(p.text)
    return "\n".join(partes)


def xml_part(ruta: Path, nombre: str) -> str:
    with zipfile.ZipFile(ruta) as zf:
        try:
            return zf.read(nombre).decode("utf-8", errors="replace")
        except KeyError:
            return ""


def validar_tipografia(doc: Document, ruta: Path):
    errores = []
    for p in doc.paragraphs:
        for r in p.runs:
            if not r.text.strip():
                continue
            if r.font.name and r.font.name.lower() != "arial":
                errores.append(f"run con fuente {r.font.name!r}: {r.text[:40]!r}")
            if r.font.size and abs(r.font.size.pt - 11.0) > 0.01:
                errores.append(f"run con tamaño {r.font.size.pt}: {r.text[:40]!r}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if not r.text.strip():
                            continue
                        if r.font.name and r.font.name.lower() != "arial":
                            errores.append(f"tabla con fuente {r.font.name!r}: {r.text[:40]!r}")
                        if r.font.size and abs(r.font.size.pt - 11.0) > 0.01:
                            errores.append(f"tabla con tamaño {r.font.size.pt}: {r.text[:40]!r}")
    xml = xml_part(ruta, "word/document.xml")
    assert 'w:val="21"' not in xml, f"{ruta.name}: persiste tamaño 10.5 pt (w:val=21)"
    if errores:
        raise AssertionError(f"{ruta.name}: formato distinto de Arial 11: {errores[:5]}")


def validar_ooxml(ruta: Path):
    document_xml = xml_part(ruta, "word/document.xml")
    settings_xml = xml_part(ruta, "word/settings.xml")
    footers = "\n".join(
        xml_part(ruta, name)
        for name in ["word/footer1.xml", "word/footer2.xml", "word/footer3.xml"]
    )

    assert "[[TOC]]" not in document_xml, f"{ruta.name}: índice no materializado"
    assert re.search(r"TOC|TOC \\o|w:docPartGallery", document_xml, re.I), (
        f"{ruta.name}: no se detecta índice real en OOXML"
    )
    assert "mirrorMargins" in settings_xml, f"{ruta.name}: falta mirrorMargins para doble cara"
    assert "updateFields" in settings_xml, f"{ruta.name}: falta updateFields"
    assert "PAGE" in footers.upper(), f"{ruta.name}: no se detecta campo de número de página"
    assert re.search(r'w:jc[^>]+w:val="right"', footers), f"{ruta.name}: pie no está alineado a la derecha"


def validar_estructura(ruta: Path, texto: str):
    nombre = ruta.name
    if nombre.startswith("Producto_01_"):
        requeridos = PLAN_HEADINGS
    elif nombre.startswith("Producto_07_") or nombre == REPORT.name:
        requeridos = FINAL_HEADINGS
    else:
        requeridos = PRODUCT_REPORT_HEADINGS

    for token in COMMON + requeridos:
        assert token.lower() in texto.lower(), f"{nombre}: falta sección/leyenda requerida: {token}"

    assert "Índice de Tablas o Cuadros" in texto, f"{nombre}: falta índice de tablas/cuadros"
    if nombre not in {"Producto_01_Plan_de_Trabajo.docx"}:
        assert "Índice de Gráficos" in texto, f"{nombre}: falta índice de gráficos"

    assert "Nombre de la consultoría" in texto, f"{nombre}: portada sin nombre de consultoría"
    assert "Consultor" in texto and "Noé Calle" in texto, f"{nombre}: portada sin nombre del consultor"
    assert "agosto de 2026" in texto.lower(), f"{nombre}: portada sin fecha"


def validar_contenido_especifico(nombre: str, texto: str):
    if nombre.startswith("Producto_03_"):
        for token in [
            "Apache Spark MLlib",
            "RandomForestClassificationModel",
            "AUC-PR",
            "PySpark",
            "benchmark sklearn",
        ]:
            assert token.lower() in texto.lower(), f"{nombre}: falta evidencia de modelo/rol: {token}"
        assert "modelo servido" in texto.lower(), f"{nombre}: no identifica el modelo servido"

    if nombre.startswith("Producto_06_"):
        for token in [
            "Apache Spark MLlib",
            "KMeans",
            "AUC-PR",
            "holdout",
            "Evaluación holdout KMeans Spark",
            "Benchmark Isolation Forest complementario",
        ]:
            assert token.lower() in texto.lower(), f"{nombre}: falta evidencia Spark/TDR: {token}"
        assert "no es el serving activo" in texto.lower(), (
            f"{nombre}: Isolation Forest no quedó inequívocamente como benchmark"
        )

    if nombre.startswith("Producto_07_"):
        for token in [
            "Plan de monitoreo y mantenimiento",
            "pruebas de integración",
            "Transferencia de conocimiento",
            "marcha blanca",
            "certificación",
            "RandomForestClassificationModel",
            "StandardScaler + KMeans",
        ]:
            assert token.lower() in texto.lower(), f"{nombre}: falta contenido de cierre: {token}"

    if nombre == REPORT.name:
        for token in [
            "47,254",
            "AUC-PR",
            "Spark MLlib",
            "GraphFrames",
            "dependencias institucionales",
            "Modelos operacionales y métricas holdout",
            "benchmarks metodológicos",
            "promoción explícita",
        ]:
            assert token.lower() in texto.lower(), f"{nombre}: falta evidencia consolidada: {token}"


def main() -> int:
    assert REPORT.exists(), f"Falta {REPORT}"
    assert len(PRODUCTS) == 7, f"Se esperaban 7 productos; encontrados: {len(PRODUCTS)}"

    for ruta in FILES:
        doc = Document(ruta)
        texto = all_text(doc)
        low = texto.lower()
        for token in BANNED:
            assert token.lower() not in low, f"{ruta.name}: contiene texto obsoleto/interno: {token}"
        validar_tipografia(doc, ruta)
        validar_ooxml(ruta)
        validar_estructura(ruta, texto)
        validar_contenido_especifico(ruta.name, texto)

    print(f"Documentación formal validada: {len(FILES)} DOCX")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
